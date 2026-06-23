"""
Analysis API endpoint.
POST /api/analyze - Analyze text and return prediction with statistics.
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
import io
import docx
import traceback
from pypdf import PdfReader

from app.database import get_db
from app.db_models import Prediction
from app.models import AnalyzeRequest, AnalyzeResponse, ContentStatistics, AIIndicator, FeatureImportance
from app.ml.predict import predict
from app.ml.train import is_trained

router = APIRouter(prefix="/api", tags=["analysis"])


@router.post("/analyze", response_model=AnalyzeResponse)
async def analyze_text(request: AnalyzeRequest, db: Session = Depends(get_db)):
    """
    Analyze academic text and predict AI vs Human origin.
    """
    try:
        if not is_trained():
            raise HTTPException(
                status_code=503,
                detail="Model not trained yet. Please train the model first via the admin panel."
            )

        text = request.text.strip()
        if len(text) < 20:
            raise HTTPException(
                status_code=400,
                detail="Text must be at least 20 characters long for meaningful analysis."
            )

        try:
            result = predict(text)
        except Exception as e:
            tb = traceback.format_exc()
            raise HTTPException(status_code=500, detail=f"Prediction failed: {str(e)}\n{tb}")

        # Save to database
        try:
            prediction_record = Prediction(
                text_snippet=text[:200],
                full_text=text,
                prediction=result["prediction"],
                ai_probability=result["ai_probability"],
                human_probability=result["human_probability"],
                word_count=result["statistics"]["total_words"],
                sentence_count=result["statistics"]["total_sentences"],
                avg_sentence_length=result["statistics"]["avg_sentence_length"],
                vocabulary_richness=result["statistics"]["vocabulary_richness"],
            )
            db.add(prediction_record)
            db.commit()
        except Exception as e:
            db.rollback()
            tb = traceback.format_exc()
            raise HTTPException(status_code=500, detail=f"Database save failed: {str(e)}\n{tb}")

        # Build response
        return AnalyzeResponse(
            prediction=result["prediction"],
            ai_probability=result["ai_probability"],
            human_probability=result["human_probability"],
            statistics=ContentStatistics(**result["statistics"]),
            indicators=[AIIndicator(**ind) for ind in result["indicators"]],
            feature_importances=[FeatureImportance(**fi) for fi in result["feature_importances"]],
        )
    except HTTPException as he:
        raise he
    except Exception as e:
        tb = traceback.format_exc()
        raise HTTPException(status_code=500, detail=f"Unexpected error in /analyze: {str(e)}\n{tb}")


@router.post("/analyze/file", response_model=AnalyzeResponse)
async def analyze_file(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """
    Analyze academic text uploaded via PDF, DOCX, or TXT file.
    """
    try:
        if not is_trained():
            raise HTTPException(
                status_code=503,
                detail="Model not trained yet. Please train the model first via the admin panel."
            )

        filename = file.filename or ""
        ext = filename.split(".")[-1].lower() if "." in filename else ""

        text = ""
        try:
            file_bytes = await file.read()
            if ext == "txt":
                try:
                    text = file_bytes.decode("utf-8")
                except UnicodeDecodeError:
                    text = file_bytes.decode("latin-1")
            elif ext == "docx":
                doc = docx.Document(io.BytesIO(file_bytes))
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                
                # Also extract from tables
                table_text = []
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                table_text.append(cell.text)
                
                text = "\n".join(paragraphs + table_text)
            elif ext == "pdf":
                reader = PdfReader(io.BytesIO(file_bytes))
                pages_text = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                text = "\n".join(pages_text)
            elif ext == "doc":
                raise HTTPException(
                    status_code=400,
                    detail="Legacy Word files (.doc) are not supported. Please save as modern .docx format."
                )
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file format (.{ext or 'unknown'}). Supported formats are: .pdf, .docx, .txt"
                )
        except HTTPException as he:
            raise he
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Failed to parse document: {str(e)}"
            )

        text = text.strip()
        if len(text) < 20:
            raise HTTPException(
                status_code=400,
                detail="Extracted text is too short. At least 20 characters of readable text are required."
            )

        try:
            result = predict(text)
        except Exception as e:
            tb = traceback.format_exc()
            raise HTTPException(status_code=500, detail=f"Prediction failed: {str(e)}\n{tb}")

        # Save to database
        try:
            prediction_record = Prediction(
                text_snippet=text[:200],
                full_text=text,
                prediction=result["prediction"],
                ai_probability=result["ai_probability"],
                human_probability=result["human_probability"],
                word_count=result["statistics"]["total_words"],
                sentence_count=result["statistics"]["total_sentences"],
                avg_sentence_length=result["statistics"]["avg_sentence_length"],
                vocabulary_richness=result["statistics"]["vocabulary_richness"],
            )
            db.add(prediction_record)
            db.commit()
        except Exception as e:
            db.rollback()
            tb = traceback.format_exc()
            raise HTTPException(status_code=500, detail=f"Database save failed: {str(e)}\n{tb}")

        # Build response
        return AnalyzeResponse(
            prediction=result["prediction"],
            ai_probability=result["ai_probability"],
            human_probability=result["human_probability"],
            statistics=ContentStatistics(**result["statistics"]),
            indicators=[AIIndicator(**ind) for ind in result["indicators"]],
            feature_importances=[FeatureImportance(**fi) for fi in result["feature_importances"]],
        )
    except HTTPException as he:
        raise he
    except Exception as e:
        tb = traceback.format_exc()
        raise HTTPException(status_code=500, detail=f"Unexpected error in /analyze/file: {str(e)}\n{tb}")

